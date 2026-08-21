"""
Ingestion pipeline:

    Resume/JD/Reference file
              |
              v
        Parse (PyMuPDF / python-docx)
              |
              v
        Chunk (LlamaIndex SentenceSplitter)
              |
              v
        Embed (Sentence Transformers) ----> Pinecone (semantic/vector index)
              |
              +-----------------------> BM25 index (keyword index, rag/bm25_index.py)

Both indexes are populated from the SAME chunks in one pass, so retrieval
(rag/hybrid_retriever.py) can search both and combine the results -- that's
"hybrid search."
"""
from __future__ import annotations

import os
from typing import List, Optional

from rag.pinecone_store import get_pinecone_store
from rag.bm25_index import get_bm25_store


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #
def parse_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def parse_docx(file_path: str) -> str:
    import docx  # python-docx

    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def parse_document(file_path: str) -> str:
    """Dispatch by extension. Supports .pdf, .docx, .txt."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    if ext == ".docx":
        return parse_docx(file_path)
    if ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    raise ValueError(f"Unsupported file type: {ext}")


# --------------------------------------------------------------------------- #
# Chunking (LlamaIndex)
# --------------------------------------------------------------------------- #
def chunk_text(text: str, chunk_size: int = 512, chunk_overlap: int = 64) -> List[str]:
    from llama_index.core.node_parser import SentenceSplitter
    from llama_index.core.schema import Document

    splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    nodes = splitter.get_nodes_from_documents([Document(text=text)])
    return [n.get_content() for n in nodes if n.get_content().strip()]


# --------------------------------------------------------------------------- #
# End-to-end ingestion into Pinecone (semantic) + BM25 (keyword)
# --------------------------------------------------------------------------- #
def ingest_text(
    text: str,
    user_id: str,
    session_id: str,
    document_type: str,
    skill: Optional[str] = None,
    topic: Optional[str] = None,
    source: Optional[str] = None,
) -> int:
    """Chunk once, then index into BOTH stores. Returns number of chunks stored."""
    chunks = chunk_text(text)

    pinecone_store = get_pinecone_store()
    pinecone_store.upsert_chunks(
        chunks=chunks,
        user_id=user_id,
        session_id=session_id,
        document_type=document_type,
        skill=skill,
        topic=topic,
        source=source or document_type,
    )

    bm25_store = get_bm25_store()
    bm25_store.add_documents(
        chunks=chunks,
        user_id=user_id,
        session_id=session_id,
        document_type=document_type,
        source=source or document_type,
    )

    return len(chunks)


def ingest_file(
    file_path: str,
    user_id: str,
    session_id: str,
    document_type: str,
    **kwargs,
) -> int:
    text = parse_document(file_path)
    return ingest_text(
        text=text,
        user_id=user_id,
        session_id=session_id,
        document_type=document_type,
        **kwargs,
    )
